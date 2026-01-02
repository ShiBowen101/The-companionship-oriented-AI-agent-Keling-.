from docx import Document
import os
from datetime import datetime

# 基础配置
output_dir = "./访谈记录"  # 输出目录
total_files = 100 # 测试生成23个文档
prefix = "访谈"  # 文件名前缀

# 自定义时间配置（需满足YYYYMMDD格式）
custom_dates = [
    "20250312",
    "20250313",
    "20250314",
    "20250315",
    "20250316",
    "20250317",
    "20250318",
    "20250319",
    "20250320",
    "20250321",
    "20250322",
    "20250323",
]


# ================== 校验模块 ==================
def validate_dates(date_list):
    """验证自定义时间格式有效性"""
    for date_str in date_list:
        try:
            datetime.strptime(date_str, "%Y%m%d")
        except ValueError:
            raise ValueError(f"无效日期格式: {date_str}，请使用YYYYMMDD格式")
    print("时间格式校验通过")


validate_dates(custom_dates)  # 执行校验

# ================== 生成模块 ==================
os.makedirs(output_dir, exist_ok=True)

for i in range(1, total_files + 1):
    # 计算时间分组索引
    group_index = (i - 1) // 10

    # 时间备用策略：超过定义数量时使用最后一个时间
    if group_index >= len(custom_dates):
        group_index = len(custom_dates) - 1
        print(f"警告：自定义时间不足，文档{i}使用最后定义时间")

    timestamp = custom_dates[group_index]

    # 创建文档对象
    doc = Document()

    # 标题样式设置（居中+加粗）
    header = doc.add_heading(f'{prefix}{i} 访谈记录', level=0)
    header.alignment = 1  # 居中[4](@ref)

    # 时间段落（含加粗效果）
    time_para = doc.add_paragraph('访谈时间：')
    time_para.add_run(timestamp).bold = True  # 加粗时间戳[5](@ref)

    # 固定内容模板
    doc.add_paragraph('访谈对象：________________')
    doc.add_paragraph('主要内容：\n1. \n2. \n3. ')

    # 结构化表格（2行3列）
    table = doc.add_table(rows=2, cols=3)
    table.style = 'LightShading-Accent1'  # 表格样式[4](@ref)
    header_cells = table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '详情'
    header_cells[2].text = '备注'

    # 生成文件名（含分组时间戳）
    filename = f"{prefix}{i}_{timestamp}.docx"
    doc.save(os.path.join(output_dir, filename))

print(f"成功生成 {total_files} 个访谈文档至 {output_dir} 目录")
